from __future__ import annotations
import sys
import os
myDir = os.getcwd()
sys.path.append(myDir)
import docx
from docx import Document
from openpyxl import *
from datetime import date
import shutil
from docx.shared import Inches
import PyQt5
from PyQt5 import QtCore, QtGui, QtWidgets, uic
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QTableWidgetItem
import time
import pyautogui
from selenium import webdriver
from win32com.client import Dispatch
import mysql.connector
from mysql.connector import errorcode
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QTableWidgetItem

class funcs:
    def TipoPessoa(tp_pessoa):
        global docprocuracao
        if tp_pessoa == 'Fisica'():
            docprocuracao = Document(f"{myDir}/Forms/Procuracao_pf.docx")
        elif tp_pessoa == 'Juridica'():
            docprocuracao = Document(f"{myDir}/Forms/Procuracao_pj.docx")
    def Criar(
    tp_pessoa,
    tituc,
    uc,
    cnc,
    rg,
    cpfcnpj,
    rua,
    num,
    comp,
    bairro,
    cid,
    cep,
    disjuntor,
    cla,
    qtdmod,
    fabmod,
    modmod,
    areaarranjo,
    qtdinversor,
    fabinversor,
    modinversor,
    somamod,
    somainversor,
    lttd,
    lgtd,
    cargainsta,
    cargainstagera,
    tensao,
    dataprev,
    tpatvd,
    disjuntor2,
    ptnmod,
    qtdarranjo,
    seriearranjo,
    ptninversor,
    qtdentinversor,
    espcc,
    espca,
    lineempresa,
    linecnpj):
        wb = load_workbook(f"{myDir}/Forms/formularioEXL.xlsx")
        docmemorial = Document(f"{myDir}/Forms/Memorial.docx")
        wb2 = load_workbook(f"{myDir}/Forms/DadosAneel.xlsx")
        home_directory = os.path.expanduser( '~' )
        tables = docmemorial.tables
        cord = (lttd + lgtd)
        data = date.today()
        curdate = data.strftime('%d/%m/%Y')
        global fase

        if disjuntor == "63A":
            inv = "63A.jpg"
            inv2 = "63A.dwg"
        if disjuntor == "100A":
            inv = "100A.jpg"
            inv2 = "100A.dwg"
        if disjuntor =="70A":
            inv = "70A.jpg"
            inv2 = "70A.dwg"

    def CriarFormularios (
            tp_pessoa,
            tituc,
            uc,
            cnc,
            rg,
            cpfcnpj,
            rua,
            num,
            comp,
            bairro,
            cid,
            cep,
            disjuntor,
            cla,
            qtdmod,
            fabmod,
            modmod,
            areaarranjo,
            qtdinversor,
            fabinversor,
            modinversor,
            somamod,
            somainversor,
            lttd,
            lgtd,
            cargainsta,
            cargainstagera,
            tensao,
            dataprev,
            tpatvd,
            disjuntor2,
            ptnmod,
            qtdarranjo,
            seriearranjo,
            ptninversor,
            qtdentinversor,
            espcc,
            espca,
            home,
            pot,
            pot2,
            inv63a,
            inv100a,
            inv70a,
            lineempresa,
            linecnpj):
        
        '''Carregamento de formulários:'''
        #Procuração
        if tp_pessoa == 'Fisica':
            docprocuracao = Document(f"{myDir}/Forms/Procuracao_pf.docx")
        elif tp_pessoa == 'Juridica':
            docprocuracao = Document(f"{myDir}/Forms/Procuracao_pj.docx")
        #Formulario EDP
        wb = load_workbook(f"{myDir}/Forms/formularioEXL.xlsx")
        #Memorial
        docmemorial = Document(f"{myDir}/Forms/Memorial.docx")
        #Dados Aneel
        wb2 = load_workbook(f"{myDir}/Forms/DadosAneel.xlsx")

        home_directory = os.path.expanduser( '~' )
        tables = docmemorial.tables
        cord = (lttd + lgtd)
        data = date.today()
        curdate = data.strftime('%d/%m/%Y')

        if disjuntor == "63A":
            inv = "63A.jpg"
            inv2 = "63A.dwg"
        if disjuntor == "100A":
            inv = "100A.jpg"
            inv2 = "100A.dwg"
        if disjuntor =="70A":
            inv = "70A.jpg"
            inv2 = "70A.dwg"

        newf = f"{home_directory}/Desktop/{tituc}"
        os.mkdir(newf)
        scr = f"{myDir}/Diagramas/{potf}/{fase}/{inv2}"
        shutil.copy(scr, newf)
        scr2 = f"{myDir}/Diagramas/{potf}/KIT.pdf"
        shutil.copy(scr2, newf)
        scr3 = f"{myDir}/Diagramas/{potf}/DATASHEET.pdf"
        shutil.copy(scr3, newf)
        scr4 = f"{myDir}/Diagramas/{potf}/INMETRO.pdf"
        shutil.copy(scr4, newf)

        referencias = {
            "AAAAA": uc,
            "BBBBB": cla,
            "CCCCC": tituc,
            "DDDDD": rua,
            "EEEEE": num,
            "FFFFF": comp,
            "GGGGG": bairro,
            "HHHHH": cid,
            "IIIII": cep,
            "JJJJJ": cargainsta,
            "MMMMM": cpfcnpj,
            "NNNNN": disjuntor,
            "OOOOO": qtdmod,
            "PPPPP": fabmod,
            "QQQQQ": modmod,
            "SSSSS": qtdinversor,
            "TTTTT": fabinversor,
            "UUUUU": modinversor,
            "VVVVV": somamod,
            "WWWWW": somainversor,
            "XXXXX": cord,
            "ZZZZZ": ptnmod,
            "AAAAB": qtdarranjo,
            "AAAAC": seriearranjo,
            "AAAAD": areaarranjo,
            "AAAAE": disjuntor2,
            "AAAAF": espcc,
            "AAAAG": espca,
            "AAAAI": rg,
            "AAAAJ": ptninversor,
            "AAAAK": qtdentinversor,
            "AAAAL": cnc,
            "AAAAM": tpatvd,
            "AAAAN": lineempresa,
            "AAAAO": linecnpj
        }

        #Formulario10kw
        sh = wb['MICROGERAÇÃO <= 10KW']
        sh['F5'] = tituc
        sh['F4'] = uc
        sh['U4'] = cla
        sh['F10'] = cpfcnpj
        sh['F6'] = rua
        sh['AC6'] = cep
        sh['U7'] = cid
        sh['X6'] = num
        sh['F7'] = bairro
        sh['R12'] = lttd
        sh['AB12'] = lgtd
        sh['I13'] = cargainsta
        sh['L16'] = cargainstagera
        sh['V13'] = tensao
        sh['M38'] = curdate

        wb.save(filename=f"Formulario{uc}.xlsx")
        scr = f"Formulario{uc}.xlsx"
        dest = f"{home_directory}/Desktop/{tituc}"
        shutil.move(scr, dest)
        #DadosAneel
        sh2 = wb2['Dados']
        sh2['C11'] = tituc
        sh2['C12'] = uc
        sh2['C13'] = tpatvd
        sh2['C15'] = disjuntor
        sh2['C17'] = cpfcnpj
        sh2['C18'] = f"{rua}, {num}, {comp}, {bairro}"
        sh2['C19'] = cep
        sh2['C20'] = cid
        sh2['C24'] = qtdmod
        sh2['C25'] = fabmod
        sh2['C26'] = modmod
        sh2['C27'] = areaarranjo
        sh2['C28'] = qtdinversor
        sh2['C29'] = fabinversor
        sh2['C30'] = modinversor
        sh2['C31'] = somamod
        sh2['C32'] = somainversor
        sh2['C33'] = dataprev
        wb2.save(filename=f"DadosAneel{uc}.xlsx")
        scr = f"DadosAneel{uc}.xlsx"
        dest = f"{home_directory}/Desktop/{tituc}"
        shutil.move(scr, dest)
        #Procuracao
        for paragrafo in docprocuracao.paragraphs:
            for codigo in referencias:
                valor = referencias[codigo]
                paragrafo.text = paragrafo.text.replace(codigo, valor)

        docprocuracao.save(f"Procuracao{uc}.docx")
        scr = f"Procuracao{uc}.docx"
        dest = f"{home_directory}/Desktop/{tituc}"
        shutil.move(scr, dest)
        #Memorial
        for paragrafo in docmemorial.paragraphs:
            for codigo in referencias:
                valor = referencias[codigo]
                paragrafo.text = paragrafo.text.replace(codigo, valor)
        #Adicionar imagens padrao
        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(f"{myDir}/image/img1.jpg" ,width=Inches(5), height=Inches(5))

        p = tables[1].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(f"{myDir}/image/img2.png" ,width=Inches(5), height=Inches(5))

        p = tables[2].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(f"{myDir}/image/img3.png" ,width=Inches(7), height=Inches(3))

        p = tables[3].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(f"{myDir}/image/img4.png" ,width=Inches(5), height=Inches(4))

        p = tables[4].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(f"{myDir}/image/img5.png" ,width=Inches(3), height=Inches(3))

        p = tables[6].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(f"{myDir}/image/img6.jpg" ,width=Inches(7), height=Inches(3))

        #Doc de identificacao
        caminho1 = QFileDialog.getOpenFileName(None, "Anexar documento")
        file1 = caminho1[0]

        p = tables[5].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(f"{file1}" ,width=Inches(5), height=Inches(5))
        #TRT
        caminho2 = QFileDialog.getOpenFileName(None, "Anexar TRT")
        file2 = caminho2[0]
            
        p = tables[7].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(f"{file2}",width=Inches(5), height=Inches(5))

        #Diagrama
        p = tables[8].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(f"{myDir}/Diagramas/{potf}/{fase}/{inv}" ,width=Inches(7), height=Inches(3))

        #Print Maps
        nav = webdriver.Chrome()
        nav.get("https://www.google.com.br/maps")
        time.sleep(4)
        nav.find_element("xpath", '//*[@id="minimap"]/div/div[2]/button').click()
        time.sleep(3)
        nav.find_element("xpath", '//*[@id="searchboxinput"]').send_keys(f"{rua} {num} {bairro} {cid} {cep}")
        time.sleep(1)
        nav.find_element("xpath", '//*[@id="searchbox-searchbutton"]').click()
        time.sleep(3)
        nav.find_element("xpath", '//*[@id="widget-zoom-in"]/div').click()
        time.sleep(1)
        nav.find_element("xpath", '//*[@id="widget-zoom-in"]/div').click()
        time.sleep(1)
        scrsht = pyautogui.screenshot(region=(500,400,400,300))
        scrsht.save("maps.png")
        nav.close()

        p = tables[12].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(f"{myDir}/maps.png" ,width=Inches(7), height=Inches(3))

        docmemorial.save(f"Memorial{uc}.docx")
        scr = f"Memorial{uc}.docx"
        dest = f"{home_directory}/Desktop/{tituc}"
        shutil.move(scr, dest)

    def CriarFormularios (
            tp_pessoa,
            tituc,
            uc,
            cnc,
            rg,
            cpfcnpj,
            rua,
            num,
            comp,
            bairro,
            cid,
            cep,
            disjuntor,
            cla,
            qtdmod,
            fabmod,
            modmod,
            areaarranjo,
            qtdinversor,
            fabinversor,
            modinversor,
            somamod,
            somainversor,
            lttd,
            lgtd,
            cargainsta,
            cargainstagera,
            tensao,
            dataprev,
            tpatvd,
            disjuntor2,
            ptnmod,
            qtdarranjo,
            seriearranjo,
            ptninversor,
            qtdentinversor,
            espcc,
            espca,
            home,
            pot,
            pot2,
            inv63a,
            inv100a,
            inv70a,
            lineempresa,
            linecnpj):
        
        '''Carregamento de formulários:'''
        #Procuração
        if tp_pessoa == 'Fisica':
            docprocuracao = Document(f"{myDir}/Forms/Procuracao_pf.docx")
        elif tp_pessoa == 'Juridica':
            docprocuracao = Document(f"{myDir}/Forms/Procuracao_pj.docx")
        #Formulario EDP
        wb = load_workbook(f"{myDir}/Forms/formularioEXL.xlsx")
        #Memorial
        docmemorial = Document(f"{myDir}/Forms/Memorial.docx")
        #Dados Aneel
        wb2 = load_workbook(f"{myDir}/Forms/DadosAneel.xlsx")

        home_directory = os.path.expanduser( '~' )
        tables = docmemorial.tables
        cord = (lttd + lgtd)
        data = date.today()
        curdate = data.strftime('%d/%m/%Y')

        if disjuntor == "63A":
            inv = "63A.jpg"
            inv2 = "63A.dwg"
        if disjuntor == "100A":
            inv = "100A.jpg"
            inv2 = "100A.dwg"
        if disjuntor =="70A":
            inv = "70A.jpg"
            inv2 = "70A.dwg"

        newf = f"{home_directory}/Desktop/{tituc}"
        os.mkdir(newf)
        scr = f"{myDir}/Diagramas/{potf}/{fase}/{inv2}"
        shutil.copy(scr, newf)
        scr2 = f"{myDir}/Diagramas/{potf}/KIT.pdf"
        shutil.copy(scr2, newf)
        scr3 = f"{myDir}/Diagramas/{potf}/DATASHEET.pdf"
        shutil.copy(scr3, newf)
        scr4 = f"{myDir}/Diagramas/{potf}/INMETRO.pdf"
        shutil.copy(scr4, newf)

        referencias = {
            "AAAAA": uc,
            "BBBBB": cla,
            "CCCCC": tituc,
            "DDDDD": rua,
            "EEEEE": num,
            "FFFFF": comp,
            "GGGGG": bairro,
            "HHHHH": cid,
            "IIIII": cep,
            "JJJJJ": cargainsta,
            "MMMMM": cpfcnpj,
            "NNNNN": disjuntor,
            "OOOOO": qtdmod,
            "PPPPP": fabmod,
            "QQQQQ": modmod,
            "SSSSS": qtdinversor,
            "TTTTT": fabinversor,
            "UUUUU": modinversor,
            "VVVVV": somamod,
            "WWWWW": somainversor,
            "XXXXX": cord,
            "ZZZZZ": ptnmod,
            "AAAAB": qtdarranjo,
            "AAAAC": seriearranjo,
            "AAAAD": areaarranjo,
            "AAAAE": disjuntor2,
            "AAAAF": espcc,
            "AAAAG": espca,
            "AAAAI": rg,
            "AAAAJ": ptninversor,
            "AAAAK": qtdentinversor,
            "AAAAL": cnc,
            "AAAAM": tpatvd,
            "AAAAN": lineempresa,
            "AAAAO": linecnpj
        }