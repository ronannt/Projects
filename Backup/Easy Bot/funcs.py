from __future__ import annotations
import sys
import os
import docx
from docx import Document
from openpyxl import *
from datetime import date
import shutil
from docx.shared import Inches
import PyQt5
from PyQt5 import QtCore, QtGui, QtWidgets, uic
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QTableWidgetItem
import time
import pyautogui
from selenium import webdriver
from win32com.client import Dispatch
import mysql.connector
from mysql.connector import errorcode
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QTableWidgetItem
myDir = os.getcwd()
sys.path.append(myDir)
home_directory = os.path.expanduser( '~' )


class functions:
    def ProcuracaoPF(nome, cpf, rg, rua, numero, bairro, cidade, estado, cep, cnc, uc):
        referencias = {
            "AAAAA": nome,
            "BBBBB": cpf,
            "CCCCC": rg,
            "DDDDD": rua,
            "EEEEE": numero,
            "FFFFF": bairro,
            "GGGGG": cidade,
            "HHHHH": estado,
            "IIIII": cep,
            "JJJJJ": cnc,
            "KKKKK": uc
            }
        docprocuracao = Document(f"{myDir}/Forms/Procuracao_pf.docx")
        for paragrafo in docprocuracao.paragraphs:
            for codigo in referencias:
                valor = referencias[codigo]
                paragrafo.text = paragrafo.text.replace(codigo, valor)

        newf = f"{home_directory}/Desktop/{nome}"
        os.mkdir(newf)
        docprocuracao.save(f"Procuracao{uc}.docx")
        scr = f"Procuracao{uc}.docx"
        dest = f"{home_directory}/Desktop/{nome}"
        shutil.move(scr, dest)

    def ProcuracaoPJ(nome, cpf, rg, rua, numero, bairro, cidade, estado, cep, cnc, uc, empresa, cnpj):
        referencias = {
            "AAAAA": nome,
            "BBBBB": cpf,
            "CCCCC": rg,
            "DDDDD": rua,
            "EEEEE": numero,
            "FFFFF": bairro,
            "GGGGG": cidade,
            "HHHHH": estado,
            "IIIII": cep,
            "JJJJJ": cnc,
            "KKKKK": uc,
            "LLLLL": empresa,
            "MMMMM": cnpj
            }
        docprocuracao = Document(f"{myDir}/Forms/Procuracao_pj.docx")
        for paragrafo in docprocuracao.paragraphs:
            for codigo in referencias:
                valor = referencias[codigo]
                paragrafo.text = paragrafo.text.replace(codigo, valor)

        newf = f"{home_directory}/Desktop/{nome}"
        os.mkdir(newf)
        docprocuracao.save(f"Procuracao{uc}.docx")
        scr = f"Procuracao{uc}.docx"
        dest = f"{home_directory}/Desktop/{nome}"
        shutil.move(scr, dest)
    
    def DadosAneel(nome, cpf, rua, numero, bairro, cidade, cep, uc, classe, disjuntor, qtdmod, fabmod, modmod, areaarranjo, qtdinversor, fabinversor, modinversor, somamod, somainversor, dataprev):
        wb2 = load_workbook(f"{myDir}/Forms/DadosAneel.xlsx")
        sh2 = wb2['Dados']
        sh2['C11'] = nome
        sh2['C12'] = uc
        sh2['C13'] = classe
        sh2['C15'] = disjuntor
        sh2['C17'] = cpf
        sh2['C18'] = f"{rua}, {numero}, {bairro}"
        sh2['C19'] = cep
        sh2['C20'] = cidade
        sh2['C24'] = qtdmod
        sh2['C25'] = fabmod
        sh2['C26'] = modmod
        sh2['C27'] = areaarranjo
        sh2['C28'] = qtdinversor
        sh2['C29'] = fabinversor
        sh2['C30'] = modinversor
        sh2['C31'] = somamod
        sh2['C32'] = somainversor
        sh2['C33'] = dataprev
        newf = f"{home_directory}/Desktop/{nome}"
        os.mkdir(newf)
        wb2.save(filename=f"DadosAneel{uc}.xlsx")
        scr = f"DadosAneel{uc}.xlsx"
        dest = f"{home_directory}/Desktop/{nome}"
        shutil.move(scr, dest)

    def Formulario10KW(nome, cpf, rua, numero, bairro, cidade, cep, uc, classe, lttd, lgtd, cargainsta, cargainstagera, tensao):
        data = date.today()
        curdate = data.strftime('%d/%m/%Y')
        wb = load_workbook(f"{myDir}/Forms/formularioEXL.xlsx")
        sh = wb['MICROGERAÇÃO <= 10KW']
        sh['F5'] = nome
        sh['F4'] = uc
        sh['U4'] = classe
        sh['F10'] = cpf
        sh['F6'] = rua
        sh['AC6'] = cep
        sh['U7'] = cidade
        sh['X6'] = numero
        sh['F7'] = bairro
        sh['R12'] = lttd
        sh['AB12'] = lgtd
        sh['I13'] = cargainsta
        sh['L16'] = cargainstagera
        sh['V13'] = tensao
        sh['M38'] = curdate

        newf = f"{home_directory}/Desktop/{nome}"
        os.mkdir(newf)
        wb.save(filename=f"Formulario{uc}.xlsx")
        scr = f"Formulario{uc}.xlsx"
        dest = f"{home_directory}/Desktop/{nome}"
        shutil.move(scr, dest)
        